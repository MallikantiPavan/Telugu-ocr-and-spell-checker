from fastapi import FastAPI, UploadFile, File, Form
from fastapi.responses import HTMLResponse, FileResponse
from PIL import Image
from ocr import extract_text, test
from spellcheck import spell_check_text
import io
import uvicorn
from docx import Document
import os
import uuid

app = FastAPI()

@app.get("/", response_class=HTMLResponse)
async def index():
    with open("static/index.html", "r", encoding="utf-8") as f:
        return f.read()

@app.post("/ocr")
async def ocr(file: UploadFile = File(...), lang: str = Form("tel")):
    try:
        image = Image.open(io.BytesIO(await file.read()))
    except Exception:
        return {"error": "Invalid image file"}

    text = extract_text(image, lang=lang)
    stats = test(text)

    return {"text": text, "stats": stats}

@app.post("/ocr-docx")
async def ocr_to_docx(file: UploadFile = File(...), lang: str = Form("tel")):
    try:
        image = Image.open(io.BytesIO(await file.read()))
    except Exception:
        return {"error": "Invalid image file"}

    text = extract_text(image, lang=lang)
    stats = test(text)

    doc = Document()
    doc.add_heading("OCR Result", level=1)

    doc.add_heading("Extracted Text", level=2)
    doc.add_paragraph(text)

    doc.add_heading("Statistics", level=2)
    for k, v in stats.items():
        doc.add_paragraph(f"{k}: {v}")

    filename = f"ocr_result_{uuid.uuid4().hex}.docx"
    filepath = os.path.join("static", filename)
    doc.save(filepath)

    return FileResponse(
        path=filepath,
        filename="ocr_result.docx",
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )


@app.post("/spell-check")
async def spell_check(text: str = Form(...), lexicon: str = Form(""), lang: str = Form("eng")):
    custom_lexicon = [line.strip() for line in lexicon.splitlines() if line.strip()]
    result = spell_check_text(text=text, custom_lexicon=custom_lexicon, language=lang)
    return result

if __name__ == "__main__":
    uvicorn.run(app, host="0.0.0.0", port=8000)
